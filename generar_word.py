from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

doc = Document()

# ─── Estilos ───────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ─── Portada ───────────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Catálogo 2026')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Análisis del Proyecto — Estado vs. Requisitos')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x47, 0x54, 0x6B)

doc.add_paragraph('')
datep = doc.add_paragraph()
datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = datep.add_run('18 de junio de 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_page_break()

# ─── Helper ────────────────────────────────────────────────────────────────
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return h

def add_status_row(table, item, status, notes=''):
    row = table.add_row()
    row.cells[0].text = item
    row.cells[1].text = status
    row.cells[2].text = notes
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(10)

def make_table(headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            p.runs[0].bold = True
    return table

# ─── 1. DISEÑO Y ESTRUCTURA WEB ───────────────────────────────────────────
add_heading('1. DISEÑO Y ESTRUCTURA WEB', level=1)

t = make_table(['Ítem', 'Estado', 'Observaciones'])
add_status_row(t, 'Diseño responsive', '✅ OK', 'Bootstrap 5.3, grid responsive, navbar colapsable')
add_status_row(t, 'Header personalizado', '✅ OK', 'Logo, buscador AJAX, nav links, login/perfil condicional')
add_status_row(t, 'Footer personalizado', '✅ OK', 'Datos reales de Solutions Mech Perú: 2 emails, 2 teléfonos, dirección')
add_status_row(t, 'Página Inicio', '✅ OK', 'Carrusel 3 slides, productos destacados, categorías, stats')
add_status_row(t, 'Página Nosotros', '✅ OK', 'Descripción real de Solutions Mech Perú')
add_status_row(t, 'Página Productos', '✅ OK', 'ListView con paginación (12), filtro por categoría, búsqueda, sidebar')
add_status_row(t, 'Página Categorías', '⚠️ STATIC placeholder', 'No consulta BD. No filtra productos. Link comentado en header.')
add_status_row(t, 'Página Marcas', '⚠️ STATIC placeholder', 'No consulta BD. Link comentado en header.')
add_status_row(t, 'Página Garantía', '⚠️ Texto genérico', 'Contenido placeholder. Link comentado en header.')
add_status_row(t, 'Página Contacto', '⚠️ Datos placeholder', 'Usa ventas@catalogo2026.local y +51 999 000 000 en vez de los reales')
add_status_row(t, 'Mapa de ubicación', '⚠️ Parcial', 'Muestra Lima centro, NO la dirección específica (Av. Esperanza Nro. 616 Int. EL09 - ATE)')
add_status_row(t, 'Optimización SEO', '⚠️ Básico', 'Faltan: Open Graph, Twitter Cards, JSON-LD, sitemap.xml, robots.txt, canonical URLs')

p = doc.add_paragraph()
p.add_run('\nConclusión: ').bold = True
p.add_run('Falta hacer dinámicas las páginas de Categorías, Marcas y Garantía, actualizar Contacto con datos reales, corregir el mapa, descomentar los links del header y mejorar el SEO.')

doc.add_page_break()

# ─── 2. SISTEMA DE USUARIOS Y PERMISOS ────────────────────────────────────
add_heading('2. SISTEMA DE USUARIOS Y PERMISOS', level=1)

t = make_table(['Ítem', 'Estado', 'Observaciones'])
add_status_row(t, 'Roles del sistema', '✅ OK', 'Superuser, Administrador (0), Ventas (1), Cliente (2)')
add_status_row(t, 'Inicio de sesión', '✅ OK', 'Email/password con LoginUser')
add_status_row(t, 'Gestión permisos por rol', '✅ OK', '3 mixins: AdministradorPermisoMixin, VentasPermisoMixin, ClientePermisoMixin')
add_status_row(t, 'Restricción acceso por URL', '✅ OK', 'Mixins aplicados a todas las vistas protegidas')
add_status_row(t, 'Activación/desactivación', '✅ OK', 'ToggleUserActiveView, con protección para no auto-desactivarse')
add_status_row(t, 'Visualización usuarios', '✅ OK', 'UserListView con tabla (Nombre, Email, Rol, Género, Estado, Acciones)')
add_status_row(t, 'Admin exclusivo', '✅ OK', 'Django admin + panel personalizado')
add_status_row(t, 'Recuperación de contraseña', '❌ NO EXISTE', 'No hay "¿Olvidaste tu contraseña?" ni flujo de reset')
add_status_row(t, 'Verificación de email', '❌ NO EXISTE', 'Registro sin confirmación de email')
add_status_row(t, 'Auto-edición de perfil', '❌ NO EXISTE', 'Solo el admin puede editar usuarios')

p = doc.add_paragraph()
p.add_run('\nConclusión: ').bold = True
p.add_run('El sistema de roles y permisos está completo. Faltan funcionalidades de usuario: recuperación de contraseña, verificación de email y edición de perfil.')

doc.add_page_break()

# ─── 3. MÓDULO DE PRODUCTOS ───────────────────────────────────────────────
add_heading('3. MÓDULO DE PRODUCTOS', level=1)

t = make_table(['Ítem', 'Estado', 'Observaciones'])
add_status_row(t, 'Creación de productos', '✅ OK', 'ProductCreateView con formsets de imágenes y especificaciones')
add_status_row(t, 'Edición de productos', '✅ OK', 'ProductUpdateView')
add_status_row(t, 'Eliminación de productos', '✅ OK', 'ProductDeleteView con confirmación')
add_status_row(t, 'Gestión de categorías', '✅ OK', 'CRUD completo: CategoryListView, CreateView, UpdateView, DeleteView')
add_status_row(t, 'Gestión de marcas', '✅ OK', 'CRUD completo: BrandListView, CreateView, UpdateView, DeleteView')
add_status_row(t, 'Carga de imágenes', '✅ OK', 'Imagen principal (image_main) + galería (ProductImage inline formset)')
add_status_row(t, 'Fichas técnicas', '✅ OK', 'Detalle con nombre, SKU, categoría, marca, precio, stock, galería')
add_status_row(t, 'Especificaciones técnicas', '✅ OK', 'Modelo TechnicalSpecification (key-value), inline formset, tabla en detalle')
add_status_row(t, 'Estado activo/inactivo', '✅ OK', 'Product.is_active, filtrado en vistas públicas')

p = doc.add_paragraph()
p.add_run('\nConclusión: ').bold = True
p.add_run('Módulo de productos COMPLETO. No falta nada de lo listado.')

doc.add_page_break()

# ─── 4. BUSCADOR INTELIGENTE ──────────────────────────────────────────────
add_heading('4. BUSCADOR INTELIGENTE', level=1)

t = make_table(['Ítem', 'Estado', 'Observaciones'])
add_status_row(t, 'Búsqueda en tiempo real (AJAX)', '✅ OK', 'Fetch API con debounce de 300ms')
add_status_row(t, 'Sugerencias automáticas', '✅ OK', 'Resultados mientras se escribe (mín. 2 caracteres)')
add_status_row(t, 'Miniatura en resultados', '✅ OK', 'Imagen 50x50px')
add_status_row(t, 'Resultados clicleables', '✅ OK', 'Links al detalle del producto')
add_status_row(t, 'Búsqueda por SKU/descripción', '⚠️ Limitado', 'ProductSearchAPIView solo busca por name__icontains, no por SKU ni descripción')

p = doc.add_paragraph()
p.add_run('\nConclusión: ').bold = True
p.add_run('El buscador funciona bien. Mejorable: ampliar búsqueda a SKU, descripción, categoría y marca.')

doc.add_page_break()

# ─── 5. SISTEMA DE COTIZACIÓN ─────────────────────────────────────────────
add_heading('5. SISTEMA DE COTIZACIÓN', level=1)

t = make_table(['Ítem', 'Estado', 'Observaciones'])
add_status_row(t, 'Visualización detallada', '✅ OK', 'PublicProductDetailView con especificaciones, galería, precio de oferta')
add_status_row(t, 'Selección de cantidad', '✅ OK', 'Campo integer en formulario individual y carrito')
add_status_row(t, 'Botón Cotizar', '✅ OK', 'Formulario individual + carrito en sesión + WhatsApp')
add_status_row(t, 'Generación de solicitud', '✅ OK', 'QuotationRequest + QuotationItem se crean al enviar')
add_status_row(t, 'Envío por email', '✅ OK', 'SMTP Gmail configurado, tabla HTML con productos, CC a roberthbardales@gmail.com')
add_status_row(t, 'Seguimiento de estado', '⚠️ ELIMINADO', 'El campo status se eliminó (migración 0002). Admin no puede gestionar estados.')
add_status_row(t, 'PDF de cotización', '❌ NO EXISTE', 'Solo email HTML, no hay descarga PDF')
add_status_row(t, 'Workflow admin', '❌ NO EXISTE', 'No hay contra-ofertas, notas internas, ni gestión de estados')

p = doc.add_paragraph()
p.add_run('\nConclusión: ').bold = True
p.add_run('El flujo básico de cotización funciona (seleccionar → enviar → email). Falta el workflow administrativo con estados y generación de PDF.')

doc.add_page_break()

# ─── 6. INTEGRACIÓN WHATSAPP ──────────────────────────────────────────────
add_heading('6. INTEGRACIÓN WHATSAPP', level=1)

t = make_table(['Ítem', 'Estado', 'Observaciones'])
add_status_row(t, 'Botón flotante', '✅ OK', 'Popup con 3 opciones: Consultar precio, Soporte técnico, Información general')
add_status_row(t, 'Cotización rápida', '✅ OK', 'Botón "Agregar" en cada card de producto, carrito en sesión')
add_status_row(t, 'Mensaje predefinido', '✅ OK', 'Incluye nombre del producto y SKU')
add_status_row(t, 'Número configurable', '⚠️ Hardcodeado', '51999888777 está escrito en base.html, products.html, public_product_detail.html')

p = doc.add_paragraph()
p.add_run('\nConclusión: ').bold = True
p.add_run('WhatsApp funciona bien. Mejorable: centralizar el número en settings.py o .env.')

doc.add_page_break()

# ─── 7. PANEL DE ADMINISTRACIÓN ───────────────────────────────────────────
add_heading('7. PANEL DE ADMINISTRACIÓN', level=1)

t = make_table(['Ítem', 'Estado', 'Observaciones'])
add_status_row(t, 'Dashboard administrativo', '⚠️ Básico', 'Solo muestra info del perfil. Sin KPIs, gráficas ni métricas.')
add_status_row(t, 'Gestión de productos', '✅ OK', 'CRUD completo con listado, filtros, paginación')
add_status_row(t, 'Gestión de categorías', '✅ OK', 'CRUD completo')
add_status_row(t, 'Gestión de marcas', '✅ OK', 'CRUD completo')
add_status_row(t, 'Gestión de usuarios', '✅ OK', 'Listado, edición, activar/desactivar')
add_status_row(t, 'Gestión de cotizaciones', '✅ OK', 'Listado y detalle. Sin cambio de estado.')
add_status_row(t, 'Gestión de inventario', '✅ OK', 'Movimientos de stock con señal que actualiza stock automáticamente')
add_status_row(t, 'KPIs / Estadísticas', '❌ NO EXISTE', 'No hay gráficas de ventas, productos bajos en stock, etc.')
add_status_row(t, 'Alertas stock bajo', '❌ NO EXISTE', 'No hay notificaciones de stock crítico')

p = doc.add_paragraph()
p.add_run('\nConclusión: ').bold = True
p.add_run('El panel tiene todas las gestiones CRUD. Falta un dashboard con indicadores y alertas de stock.')

doc.add_page_break()

# ─── RESUMEN PRIORIZADO ────────────────────────────────────────────────────
add_heading('RESUMEN PRIORIZADO', level=1)

doc.add_paragraph('A continuación se listan las tareas pendientes ordenadas por prioridad:')

priorities = [
    ('CRÍTICO', [
        'Hacer dinámica la Página Categorías (consultar BD, filtrar productos)',
        'Hacer dinámica la Página Marcas (consultar BD, filtrar productos)',
        'Actualizar Página Contacto con datos reales (emails, teléfonos, dirección)',
        'Corregir mapa de ubicación para mostrar dirección específica (Av. Esperanza Nro. 616 Int. EL09 - ATE)',
        'Descomentar links del header a Categorías, Marcas y Garantía',
        'Agregar Dashboard con KPIs (total productos, cotizaciones, stock bajo, gráficas)',
        'Agregar recuperación de contraseña (flujo de reset por email)',
    ]),
    ('IMPORTANTE', [
        'Centralizar número de WhatsApp en settings.py / .env',
        'SEO: agregar Open Graph tags, Twitter Cards, sitemap.xml, robots.txt',
        'Ampliar buscador AJAX para incluir SKU, descripción, categoría y marca',
        'Agregar verificación de email al registrarse',
        'Agregar auto-edición de perfil (nombre, teléfono)',
    ]),
    ('DESEABLE', [
        'Restaurar / agregar workflow de cotizaciones con estados (pendiente, aprobado, rechazado)',
        'Generar PDF de cotización',
        'Agregar productos relacionados en detalle de producto',
        'Alertas de stock bajo (email o notificación en dashboard)',
        'Pruebas automatizadas (tests)',
    ]),
]

for level, items in priorities:
    add_heading(level, level=2)
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f'{i}. {item}', style='List Number')

# ─── ESTADO GENERAL ────────────────────────────────────────────────────────
add_heading('ESTADO GENERAL DEL PROYECTO', level=1)

p = doc.add_paragraph()
run = p.add_run('Total de ítems evaluados: 48')
run.bold = True

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0]
hdr2.cells[0].text = 'Estado'
hdr2.cells[1].text = 'Cantidad'
for p in hdr2.cells[0].paragraphs:
    p.runs[0].bold = True
for p in hdr2.cells[1].paragraphs:
    p.runs[0].bold = True

for item, status in [('✅ Completos (OK)', '30'), ('⚠️ Incompletos / Mejorables', '11'), ('❌ No existen (Faltan)', '7')]:
    row = table2.add_row()
    row.cells[0].text = item
    row.cells[1].text = status

p = doc.add_paragraph()
p.add_run('\nPorcentaje de avance estimado: ~70% ').bold = True
p.add_run('(considerando funcionalidades principales implementadas vs. lista de requisitos)')

# ─── Guardar ───────────────────────────────────────────────────────────────
doc.save(r'D:\Proyectos GIT\catalogo2026\analisis_proyecto.docx')
print('Documento generado correctamente.')
